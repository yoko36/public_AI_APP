from __future__ import annotations
import io
import json
import mimetypes
import os
import re
import tempfile
from typing import Dict, Iterable, List, Optional, Tuple

import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import SupabaseVectorStore
from supabase import create_client, Client

# ==================================================
## 環境設定
# ==================================================
SUPABASE_URL = os.environ["SUPABASE_URL"]
SUPABASE_SERVICE_ROLE_KEY = os.environ["SUPABASE_SERVICE_ROLE_KEY"]
OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]

EMBED_MODEL = os.getenv("EMBED_MODEL", "text-embedding-3-small")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "500"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "100"))

emb = OpenAIEmbeddings(model=EMBED_MODEL, api_key=OPENAI_API_KEY)

# ==================================================
## 補助定義
# ==================================================

_TEXT_MIME = {
    "text/plain",
    "text/markdown",
    "text/x-markdown",
    "text/html",
    "application/json",
    "text/csv",
    "text/tab-separated-values",
}

_DOC_MIME = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
}

_PPT_MIME = {
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",  # .pptx
}

_PDF_MIME = {"application/pdf"}

_CSV_MIME = {"text/csv", "text/tab-separated-values", "application/vnd.ms-excel"}

# 拡張子 -> MIME の対応を拡張（マークダウンとかはmimetypesに登録されていない）
mimetypes.add_type("text/markdown", ".md")
mimetypes.add_type("text/markdown", ".markdown")
mimetypes.add_type("text/tab-separated-values", ".tsv")


# MIMEタイプの識別
def _guess_mime(object_path: str, provided: Optional[str]) -> str:
    # MIME を推定する（提供されていればそれを優先）
    if provided:
        return provided
    mime, _ = mimetypes.guess_type(object_path)
    return mime or "application/octet-stream"


# ファイル内のPostgreSQLで使用できない文字を省くための正規表現
_CTRL_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def _clean_text(s: str) -> str:
    """
    Postgres/JSON が嫌う NUL(\x00) を含む制御文字を空白に置換。
    22P05 (\u0000 cannot be converted to text) 対策。
    """
    if not s:
        return s
    return _CTRL_RE.sub(" ", s)


def _decode_bytes(data: bytes) -> str:
    # バイト列を文字列にデコードする。まず UTF-8、失敗時は推定にフォールバック
    try:
        return _clean_text(data.decode("utf-8"))
    except UnicodeDecodeError:
        # 任意: chardet があれば使用してエンコーディングを推定
        try:
            import chardet  # type: ignore

            enc = chardet.detect(data).get("encoding") or "utf-8"
            return _clean_text(data.decode(enc, errors="replace"))
        except Exception:
            return _clean_text(data.decode("utf-8", errors="replace"))


# ==================================================
## 各形式の抽出器（抽出後のデータは（”<テキスト>”, "<ページ番号> or <None>"）となる
# ==================================================


# PDFのページごとにデータを抽出（データ型は("テキスト", "ページ番号")となる）
def _extract_pdf(data: bytes) -> List[Tuple[str, int]]:
    out: List[Tuple[str, int]] = []  # 出力データ
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i, page in enumerate(doc, 1):
            text = page.get_text("text") or ""
            text = _clean_text(text)
            if text.strip():
                out.append((text, i))
    return out


# Markdownファイルの先頭にあるYAMLフロントマター（メタデータ部）を削除
def _strip_frontmatter(md: str) -> str:
    return re.sub(r"^---.*?---", "", md, flags=re.DOTALL)


# Markdown ファイルのデータ抽出（データ型は("テキスト全文", "None")となる）
def _extract_markdown(data: bytes) -> List[Tuple[str, Optional[int]]]:
    text = _strip_frontmatter(_decode_bytes(data))
    return [(text, None)]


# 平文の抽出
def _extract_plain(data: bytes) -> List[Tuple[str, Optional[int]]]:
    return [(_decode_bytes(data), None)]


# HTMLの解析
def _has_lxml() -> bool:
    # lxml が利用可能かどうかを返す
    try:
        import lxml  # type: ignore  # noqa: F401

        return True
    except Exception:
        return False


# HTMLの抽出（不要なタグを削除したテキスト）
def _extract_html(data: bytes) -> List[Tuple[str, Optional[int]]]:
    html = _decode_bytes(data)
    soup = BeautifulSoup(html, "lxml" if _has_lxml() else "html.parser")
    # script/style/noscript を除去
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text("")
    text = re.sub(r"{2,}", "", text)  # HTMLのコメントなどを取り除く
    return [(text.strip(), None)]


# JSONの抽出
def _extract_json(data: bytes) -> List[Tuple[str, Optional[int]]]:
    # JSON を整形文字列化。失敗時はプレーンテキストとして扱う
    try:
        obj = json.loads(_decode_bytes(data))
        pretty = json.dumps(obj, ensure_ascii=False, indent=2)
        return [(pretty, None)]
    except Exception:
        return [(_decode_bytes(data), None)]


# CSV/TSVの抽出器
def _extract_csv_like(data: bytes) -> List[Tuple[str, Optional[int]]]:
    # CSV/TSV は区切りテキストとしてそのまま扱う（RAG では行単位が有用なことが多い）
    return [(_decode_bytes(data), None)]


# docx（ワードファイル）の抽出
def _extract_docx(data: bytes) -> List[Tuple[str, Optional[int]]]:
    # docx から本文と表を抽出して結合
    try:
        from docx import Document as DocxDocument  # python-docx

        with io.BytesIO(data) as bio:
            doc = DocxDocument(bio)
        texts: List[str] = []
        # word内の段落で分割
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                texts.append(p.text)
        # Word内のすべての表を走査し、各行をタブ区切り文字列として抽出
        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = "	".join(
                    cell.text.strip() for cell in row.cells
                )  # タブでつなげる
                if row_text.strip():
                    texts.append(row_text)
        joined = "".join(texts)
        return [(_clean_text(joined), None)]
    except Exception:
        # 失敗した場合はプレーンテキストとしてフォールバック
        return [(_decode_bytes(data), None)]


# pptx（パワポファイル）の抽出（出力は (スライドのテキスト, スライド番号)となる）
def _shape_texts(shape) -> list[str]:
    texts = []
    # 1) テキストフレームを持つ通常の図形（テキストボックスやプレースホルダ）
    if getattr(shape, "has_text_frame", False) and shape.text_frame:
        # 段落単位で抽出（runsまで細かくやるなら p.runs も可）
        for p in shape.text_frame.paragraphs:
            if p.text and p.text.strip():
                texts.append(p.text)
    # 2) 表（table）内のセルテキスト
    if getattr(shape, "has_table", False) and shape.table:
        for row in shape.table.rows:
            row_text = "\t".join((cell.text or "").strip() for cell in row.cells)
            if row_text.strip():
                texts.append(row_text)
    # 3) グループ化された図形の中身を再帰的に辿る
    if shape.shape_type == 6 and hasattr(shape, "shapes"):  # MSO_SHAPE_TYPE.GROUP = 6
        for s in shape.shapes:
            texts.extend(_shape_texts(s))

    return texts


# パワポ（pptx）ファイルの抽出
def _extract_pptx(slide) -> str:
    texts = []
    for shape in slide.shapes:
        texts.extend(_shape_texts(shape))
    return "".join(texts)  # 必要なら " ".join(...) に変更


# ==================================================
## チャンク化と保存
# ==================================================


def _to_documents(
    texts: List[Tuple[str, Optional[int]]],
    project_id: str,
    attachment_id: str,
    title: str,
    source: str,
) -> List[Document]:
    """抽出したテキストをチャンク分割し、LangChain の Document に変換。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    docs: List[Document] = []
    for text, page in texts:
        if not (text and text.strip()):
            continue
        text = _clean_text(text)
        for chunk in splitter.split_text(text):
            docs.append(
                Document(
                    page_content=_clean_text(chunk),
                    metadata={
                        "project_id": project_id,
                        "attachment_id": attachment_id,
                        "title": title,
                        "source": source,
                        "page": page,
                    },
                )
            )
    return docs


# ==================================================
## 公開エントリポイント
# ==================================================
def ingest_from_storage(
    *,
    storage_bucket: str,
    object_path: str,
    project_id: str,
    attachment_id: str,
    title: str,
    mime: Optional[str] = None,
    client: Optional[Client] = None,
) -> int:
    """
    Supabase Storage 上の 1 ファイルをベクトル DB に取り込む。
    戻り値: 生成・挿入されたチャンク数
    """
    sb = client or create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
    sb.postgrest.schema("app")

    # 1) ダウンロード
    data: bytes = sb.storage.from_(storage_bucket).download(object_path)
    if not isinstance(data, (bytes, bytearray)):
        try:
            data = data["data"]  # type: ignore
        except Exception:
            raise RuntimeError("Supabase SDK のダウンロード結果が想定外です")

    # 2) MIME に応じた抽出
    mime = _guess_mime(object_path, mime)
    source = f"{storage_bucket}/{object_path}"

    extracted: List[Tuple[str, Optional[int]]]

    if mime in _PDF_MIME or object_path.lower().endswith(".pdf"):
        extracted = _extract_pdf(data)
    elif mime in _DOC_MIME or object_path.lower().endswith(".docx"):
        extracted = _extract_docx(data)
    elif mime in _PPT_MIME or object_path.lower().endswith(".pptx"):
        extracted = _extract_pptx(data)
    elif mime in _TEXT_MIME or object_path.lower().endswith(
        (".txt", ".md", ".markdown", ".html", ".htm")
    ):
        if mime.startswith("text/markdown") or object_path.lower().endswith(
            (".md", ".markdown")
        ):
            extracted = _extract_markdown(data)
        elif mime in {"text/html"} or object_path.lower().endswith((".html", ".htm")):
            extracted = _extract_html(data)
        elif mime in {"application/json"} or object_path.lower().endswith(".json"):
            extracted = _extract_json(data)
        else:
            extracted = _extract_plain(data)
    elif mime in _CSV_MIME or object_path.lower().endswith((".csv", ".tsv")):
        extracted = _extract_csv_like(data)
    else:
        # 未知形式はプレーンテキストとしてフォールバック（完全に捨てない方針）
        extracted = _extract_plain(data)

    # 3) チャンク化して Document 化
    docs = _to_documents(
        texts=extracted,
        project_id=project_id,
        attachment_id=attachment_id,
        title=title,
        source=source,
    )

    if not docs:
        return 0

    # 4) ベクトルストアに挿入
    SupabaseVectorStore.from_documents(
        documents=docs,
        embedding=emb,
        client=sb,
        table_name="lc_documents",
        query_name="match_documents",
    )

    return len(docs)
